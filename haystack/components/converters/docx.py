# SPDX-FileCopyrightText: 2022-present deepset GmbH <info@deepset.ai>
#
# SPDX-License-Identifier: Apache-2.0

import io
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

from haystack import Document, component, logging
from haystack.components.converters.utils import get_bytestream_from_source, normalize_metadata
from haystack.dataclasses import ByteStream
from haystack.lazy_imports import LazyImport

logger = logging.getLogger(__name__)

with LazyImport("Run 'pip install python-docx'") as docx_import:
    import docx
    from docx.document import Document as DocxDocument


@component
class DocxToDocument:
    """
    Converts Docx files to Documents.

    Uses `python-docx` library to convert the Docx file to a document.
    This component does not preserve page brakes in the original document.

    Usage example:
    ```python
    from haystack.components.converters.docx import DocxToDocument

    converter = DocxToDocument()
    results = converter.run(sources=["sample.docx"], meta={"date_added": datetime.now().isoformat()})
    documents = results["documents"]
    print(documents[0].content)
    # 'This is a text from the Docx file.'
    ```
    """

    def __init__(self):
        """
        Create a DocxToDocument component.
        """
        docx_import.check()

    @component.output_types(documents=List[Document])
    def run(
        self,
        sources: List[Union[str, Path, ByteStream]],
        meta: Optional[Union[Dict[str, Any], List[Dict[str, Any]]]] = None,
    ):
        """
        Converts Docx files to Documents.

        :param sources:
            List of file paths or ByteStream objects.
        :param meta:
            Optional metadata to attach to the Documents.
            This value can be either a list of dictionaries or a single dictionary.
            If it's a single dictionary, its content is added to the metadata of all produced Documents.
            If it's a list, the length of the list must match the number of sources, because the two lists will be zipped.
            If `sources` contains ByteStream objects, their `meta` will be added to the output Documents.

        :returns:
            A dictionary with the following keys:
            - `documents`: Created Documents
        """
        documents = []
        meta_list = normalize_metadata(meta=meta, sources_count=len(sources))

        for source, metadata in zip(sources, meta_list):
            try:
                bytestream = get_bytestream_from_source(source)
            except Exception as e:
                logger.warning("Could not read {source}. Skipping it. Error: {error}", source=source, error=e)
                continue
            try:
                file = docx.Document(io.BytesIO(bytestream.data))
                paragraphs = [para.text for para in file.paragraphs]
                text = "\n".join(paragraphs)
            except Exception as e:
                logger.warning(
                    "Could not read {source} and convert it to a Docx Document, skipping. Error: {error}",
                    source=source,
                    error=e,
                )
                continue

            docx_meta = self._get_docx_metadata(document=file)
            merged_metadata = {**bytestream.meta, **docx_meta, **metadata}
            document = Document(content=text, meta=merged_metadata)
            documents.append(document)

        return {"documents": documents}

    def _get_docx_metadata(self, document: DocxDocument) -> Dict[str, Union[str, int, datetime]]:
        """
        Get all relevant data from the 'core_properties' attribute from a Docx Document.

        Only add metadata fields that are not None or not empty strings.

        :param document:
            The Docx Document you want to extract metadata from

        :returns:
            A dictionary containing all the relevant fields from the 'core_properties'
        """
        docx_meta = {}
        props = [
            "author",
            "category",
            "comments",
            "content_status",
            "created",
            "identifier",
            "keywords",
            "language",
            "last_modified_by",
            "last_printed",
            "modified",
            "revision",
            "subject",
            "title",
            "version",
        ]
        for prop in props:
            if hasattr(document.core_properties, prop):
                value = getattr(document.core_properties, prop)
                if value is not None and value != "":
                    docx_meta[f"docx_{prop}"] = value
        return docx_meta
